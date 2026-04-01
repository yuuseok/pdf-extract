import logging
from docx import Document

logger = logging.getLogger(__name__)


class DocxService:
    """Word(.docx) 파일에서 text, markdown, json을 추출."""

    def extract(self, file_path: str) -> dict:
        doc = Document(file_path)
        elements = []

        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            style_name = para.style.name if para.style else ""
            heading_level = self._get_heading_level(style_name)

            if heading_level:
                elements.append({
                    "type": "heading",
                    "heading_level": heading_level,
                    "content": para.text.strip(),
                })
            else:
                elements.append({
                    "type": "paragraph",
                    "content": para.text.strip(),
                })

        for table in doc.tables:
            table_data = self._extract_table(table)
            if table_data["rows"]:
                elements.append(table_data)

        text = self._to_text(elements)
        markdown = self._to_markdown(elements)

        return {
            "text": text,
            "markdown": markdown,
            "json": elements,
            "json_raw": elements,
            "page_count": None,
        }

    def _get_heading_level(self, style_name: str) -> int | None:
        if style_name.startswith("Heading"):
            try:
                return int(style_name.replace("Heading ", "").strip())
            except ValueError:
                return 1
        return None

    def _extract_table(self, table) -> dict:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        return {"type": "table", "rows": rows}

    def _to_text(self, elements: list) -> str:
        lines = []
        for el in elements:
            if el["type"] in ("heading", "paragraph"):
                lines.append(el["content"])
            elif el["type"] == "table":
                for row in el["rows"]:
                    lines.append("\t".join(row))
                lines.append("")
        return "\n".join(lines)

    def _to_markdown(self, elements: list) -> str:
        lines = []
        for el in elements:
            if el["type"] == "heading":
                prefix = "#" * el.get("heading_level", 1)
                lines.append(f"{prefix} {el['content']}")
                lines.append("")
            elif el["type"] == "paragraph":
                lines.append(el["content"])
                lines.append("")
            elif el["type"] == "table":
                rows = el["rows"]
                if not rows:
                    continue
                # 헤더
                lines.append("| " + " | ".join(rows[0]) + " |")
                lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
                for row in rows[1:]:
                    # 셀 수 맞추기
                    padded = row + [""] * (len(rows[0]) - len(row))
                    lines.append("| " + " | ".join(padded) + " |")
                lines.append("")
        return "\n".join(lines)
