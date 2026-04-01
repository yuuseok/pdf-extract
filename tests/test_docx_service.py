import os
import tempfile
from docx import Document
from app.service.docx_service import DocxService


def _create_test_docx(path: str):
    doc = Document()
    doc.add_heading("테스트 제목", level=1)
    doc.add_paragraph("첫 번째 문단입니다.")
    doc.add_heading("하위 제목", level=2)
    doc.add_paragraph("두 번째 문단입니다.")
    table = doc.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "이름"
    table.rows[0].cells[1].text = "나이"
    table.rows[0].cells[2].text = "도시"
    table.rows[1].cells[0].text = "홍길동"
    table.rows[1].cells[1].text = "30"
    table.rows[1].cells[2].text = "서울"
    doc.save(path)


def test_docx_extract_returns_three_formats():
    svc = DocxService()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        _create_test_docx(f.name)
        result = svc.extract(f.name)
    os.unlink(f.name)

    assert "text" in result
    assert "markdown" in result
    assert "json" in result
    assert "테스트 제목" in result["text"]
    assert "# 테스트 제목" in result["markdown"]
    assert len(result["json"]) > 0


def test_docx_table_extraction():
    svc = DocxService()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        _create_test_docx(f.name)
        result = svc.extract(f.name)
    os.unlink(f.name)

    assert "홍길동" in result["text"]
    assert "홍길동" in result["markdown"]
    tables = [e for e in result["json"] if e["type"] == "table"]
    assert len(tables) == 1


def test_docx_heading_levels():
    svc = DocxService()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        _create_test_docx(f.name)
        result = svc.extract(f.name)
    os.unlink(f.name)

    assert "# 테스트 제목" in result["markdown"]
    assert "## 하위 제목" in result["markdown"]
